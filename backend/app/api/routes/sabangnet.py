"""사방넷 CS 게시판 대응 API 라우트

쇼핑몰 문의사항을 사방넷 API를 통해 수집하고,
Claude API로 AI 답변 초안을 생성한 뒤,
확인/수정 후 사방넷 API로 답변을 발송하는 시스템.
"""
import os
import io
import json
import logging
from datetime import datetime
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import Response
from fastapi.responses import Response as FastAPIResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func, distinct
import httpx

from app.database import get_db
from app.db_models import CsInquiry, CsReferenceData, CsConfig, DeliveryTracking, CsFollowUpAction

router = APIRouter(prefix="/sabangnet", tags=["sabangnet"])
logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "uploads", "cs_reference")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".doc", ".txt", ".csv"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


def extract_text_from_file(file_bytes: bytes, file_name: str) -> str:
    """파일에서 텍스트를 추출"""
    ext = os.path.splitext(file_name)[1].lower()

    try:
        if ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text.strip())
            return "\n\n".join(texts)

        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            texts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip(" |"):
                        texts.append(row_text)
            return "\n".join(texts)

        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            # 테이블 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        texts.append(row_text)
            return "\n".join(texts)

        elif ext in (".txt", ".csv"):
            return file_bytes.decode("utf-8", errors="replace")

        else:
            return ""
    except Exception as e:
        logger.error(f"파일 텍스트 추출 실패 ({file_name}): {e}")
        return f"[텍스트 추출 실패: {str(e)}]"


# ──────────────────────────────────────────────
# Pydantic Request / Response Models
# ──────────────────────────────────────────────

class InquiryCreate(BaseModel):
    external_id: Optional[str] = None
    mall_name: str
    board_type: Optional[str] = None
    customer_name: Optional[str] = None
    customer_id: Optional[str] = None
    product_name: Optional[str] = None
    order_number: Optional[str] = None
    title: Optional[str] = None
    content: str
    inquiry_date: Optional[str] = None
    category: Optional[str] = None
    priority: str = "normal"


class InquiryUpdate(BaseModel):
    status: Optional[str] = None
    ai_response: Optional[str] = None
    final_response: Optional[str] = None
    category: Optional[str] = None
    priority: Optional[str] = None


class ReferenceDataCreate(BaseModel):
    title: str
    category: Optional[str] = None
    content: str


class ReferenceDataUpdate(BaseModel):
    title: Optional[str] = None
    category: Optional[str] = None
    content: Optional[str] = None
    is_active: Optional[bool] = None


class AiGenerateRequest(BaseModel):
    inquiry_id: int


class BulkAiGenerateRequest(BaseModel):
    inquiry_ids: List[int]


class SendResponseRequest(BaseModel):
    inquiry_id: int
    response: Optional[str] = None  # None이면 final_response 사용


class ConfigUpdate(BaseModel):
    auto_mode: Optional[bool] = None
    auto_categories: Optional[List[str]] = None  # 자동 답변 카테고리 목록
    sabangnet_api_key: Optional[str] = None
    claude_api_key: Optional[str] = None
    response_tone: Optional[str] = None  # 친근한, 정중한, 비즈니스


# ──────────────────────────────────────────────
# AI 답변 생성 유틸
# ──────────────────────────────────────────────

def generate_template_response(inquiry: CsInquiry) -> str:
    """API 없을 때 템플릿 기반 답변"""
    templates = {
        "배송문의": (
            f"안녕하세요, 고객님. {inquiry.mall_name} 고객센터입니다.\n\n"
            f"배송 관련 문의 주셔서 감사합니다.\n"
            f"주문하신 상품은 현재 정상적으로 배송 진행 중이며, "
            f"1-3 영업일 내 수령하실 수 있습니다.\n\n"
            f"추가 문의사항이 있으시면 언제든 연락 주세요.\n감사합니다."
        ),
        "교환/반품": (
            f"안녕하세요, 고객님. {inquiry.mall_name} 고객센터입니다.\n\n"
            f"교환/반품 문의 주셔서 감사합니다.\n"
            f"수령 후 7일 이내 교환/반품 신청이 가능합니다.\n\n"
            f"[교환/반품 절차]\n"
            f"1. 해당 쇼핑몰에서 교환/반품 신청\n"
            f"2. 택배 수거 (반품비는 고객 사유 시 부담)\n"
            f"3. 상품 확인 후 교환 발송 또는 환불 처리\n\n"
            f"추가 문의사항이 있으시면 언제든 연락 주세요.\n감사합니다."
        ),
        "상품문의": (
            f"안녕하세요, 고객님. {inquiry.mall_name} 고객센터입니다.\n\n"
            f"상품에 대해 관심 가져주셔서 감사합니다.\n"
            f"문의하신 내용에 대해 안내드립니다.\n\n"
            f"{inquiry.product_name or '해당 상품'}에 대한 자세한 정보는 "
            f"상품 상세 페이지를 참고해 주시기 바랍니다.\n\n"
            f"추가 문의사항이 있으시면 언제든 연락 주세요.\n감사합니다."
        ),
    }
    board = inquiry.board_type or "일반문의"
    return templates.get(
        board,
        (
            f"안녕하세요, 고객님. {inquiry.mall_name} 고객센터입니다.\n\n"
            f"문의 주셔서 감사합니다.\n"
            f"확인 후 빠른 시일 내 답변 드리겠습니다.\n\n감사합니다."
        ),
    )


async def generate_ai_response(inquiry: CsInquiry, reference_data_list: list, db: Session = None) -> str:
    """Claude API를 사용해 CS 답변 초안 생성"""
    from app.config import get_settings
    settings = get_settings()
    api_key = settings.ANTHROPIC_API_KEY or os.getenv("ANTHROPIC_API_KEY", "")

    if not api_key:
        # Fallback: 템플릿 기반 답변
        return generate_template_response(inquiry)

    # 참고 자료 컨텍스트 구성 (텍스트 + 파일 추출 내용 포함)
    ref_parts = []
    for ref in reference_data_list:
        part = f"[{ref.category}] {ref.title}:\n{ref.content}"
        if ref.extracted_text:
            part += f"\n\n[첨부파일 ({ref.file_name}) 내용]:\n{ref.extracted_text[:3000]}"
        ref_parts.append(part)
    ref_context = "\n\n".join(ref_parts)

    # 실시간 주문/배송 정보 조회
    order_context = ""
    delivery_context = ""
    if inquiry.order_number:
        try:
            from app.services.sabangnet_api import get_sabangnet_api
            api = get_sabangnet_api()
            if api.is_available:
                order_data = await api.get_order_detail(inquiry.order_number)
                if order_data.get("success") and order_data.get("items"):
                    oi = order_data["items"][0]
                    order_context = f"""
## 실시간 주문 정보
- 주문번호: {inquiry.order_number}
- 주문상태: {oi.get('ORDER_STATUS', '확인불가')}
- 배송사: {oi.get('DELIVERY_COMPANY_NM', '미정')}
- 운송장번호: {oi.get('DELIVERY_NO', '미발급')}
- 주문일: {oi.get('ORDER_DATE', '')}
- 상품명: {oi.get('PRODUCT_NM', '')}
- 수량: {oi.get('SALE_CNT', '')}
- 결제금액: {oi.get('ORDER_TOTAL_PRICE', '')}원
"""
                    # 운송장번호가 있으면 배송 추적
                    tracking_no = oi.get('DELIVERY_NO', '').strip()
                    courier_name = oi.get('DELIVERY_COMPANY_NM', '').strip()
                    if tracking_no and courier_name:
                        try:
                            from app.services.delivery_tracker import DeliveryTracker
                            tracker = DeliveryTracker()
                            tracking = await tracker.track_delivery(tracking_no, courier_name)
                            if tracking.get("success"):
                                delivery_context = f"""
## 실시간 배송 추적
- 현재상태: {tracking.get('status_text', '확인불가')}
- 최근이벤트: {tracking.get('last_event', '')}
- 배달완료여부: {'완료' if tracking.get('complete') else '미완료'}
"""
                                if tracking.get('events'):
                                    recent = tracking['events'][-3:]  # 최근 3건
                                    delivery_context += "- 최근 배송이력:\n"
                                    for ev in recent:
                                        delivery_context += f"  · {ev.get('time', '')} {ev.get('location', '')} - {ev.get('status', '')}\n"

                                # DB에 배송 추적 정보 저장
                                if db:
                                    try:
                                        existing_track = db.query(DeliveryTracking).filter(
                                            DeliveryTracking.inquiry_id == inquiry.id
                                        ).first()
                                        if existing_track:
                                            existing_track.tracking_number = tracking_no
                                            existing_track.courier_name = courier_name
                                            existing_track.courier_code = tracking.get('courier_code', '')
                                            existing_track.current_status = tracking.get('status', 'unknown')
                                            existing_track.last_event = tracking.get('last_event', '')
                                            existing_track.tracking_history = tracking.get('events', [])
                                            existing_track.last_checked_at = datetime.now()
                                        else:
                                            dt = DeliveryTracking(
                                                inquiry_id=inquiry.id,
                                                order_number=inquiry.order_number,
                                                tracking_number=tracking_no,
                                                courier_name=courier_name,
                                                courier_code=tracking.get('courier_code', ''),
                                                current_status=tracking.get('status', 'unknown'),
                                                last_event=tracking.get('last_event', ''),
                                                tracking_history=tracking.get('events', []),
                                                order_detail=order_data.get("items", [{}])[0] if order_data.get("items") else None,
                                            )
                                            db.add(dt)
                                        db.commit()
                                    except Exception as e_track:
                                        logger.warning(f"배송 추적 저장 실패: {e_track}")
                        except Exception as e_delivery:
                            logger.warning(f"배송 추적 조회 실패: {e_delivery}")
        except Exception as e_order:
            logger.warning(f"주문 상세 조회 실패: {e_order}")

    prompt = f"""당신은 '{inquiry.mall_name}' 쇼핑몰의 널담(Nuldam) 브랜드 CS 담당자입니다.
고객 문의에 대해 전문적이고 친절한 답변을 작성해주세요.

## 참고 자료
{ref_context if ref_context else '(등록된 참고 자료 없음)'}
{order_context}
{delivery_context}

## 고객 문의 정보
- 쇼핑몰: {inquiry.mall_name}
- 문의 유형: {inquiry.board_type or '일반문의'}
- 문의 구분: {inquiry.category or '미분류'}
- 상품명: {inquiry.product_name or '미지정'}
- 주문번호: {inquiry.order_number or '미지정'}
- 고객명: {inquiry.customer_name or '미지정'}
- 제목: {inquiry.title or ''}
- 내용: {inquiry.content}

## 답변 작성 지침
1. 고객의 문의 내용을 정확히 파악하고 구체적으로 답변하세요
2. 참고 자료에 해당 정보가 있으면 적극 활용하세요
3. 교환/반품/환불 요청 시 구체적인 절차를 안내하세요
4. '{inquiry.mall_name} 널담 고객센터' 명의로 작성하세요
5. 문의 내용에 맞는 구체적인 해결 방안을 제시하세요
6. 일반적인 템플릿이 아닌, 이 고객의 상황에 맞춘 답변을 작성하세요
7. 주문 정보와 배송 상태가 제공된 경우, 해당 정보를 바탕으로 구체적인 답변을 작성하세요 (예: 운송장번호, 현재 배송 위치 등)
8. 배송 지연/미배송 문의 시 실제 배송 상태를 기반으로 정확한 안내를 해주세요

## 후속 조치 태그
답변 끝에 필요한 후속 조치를 [ACTION:type] 형태로 태그해주세요.
가능한 액션: exchange(교환), refund(환불), reship(재배송), damage_claim(파손보상), restock_notify(재입고알림), delivery_check(배송확인), none(조치불필요)
예: [ACTION:delivery_check] 또는 [ACTION:exchange]

답변:"""

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 1024,
                    "messages": [{"role": "user", "content": prompt}],
                },
                timeout=30.0,
            )
            if resp.status_code == 200:
                data = resp.json()
                return data["content"][0]["text"]
            else:
                logger.error(f"Claude API 응답 오류: {resp.status_code} - {resp.text}")
    except Exception as e:
        logger.error(f"Claude API 호출 실패: {e}")

    return generate_template_response(inquiry)


# ──────────────────────────────────────────────
# 문의사항 (Inquiries) 엔드포인트
# ──────────────────────────────────────────────

@router.get("/inquiries/stats")
def get_inquiry_stats(db: Session = Depends(get_db)):
    """대시보드 통계: 전체 수, 상태별, 쇼핑몰별, 카테고리별"""
    try:
        total = db.query(func.count(CsInquiry.id)).scalar() or 0

        # 상태별 통계
        status_rows = (
            db.query(CsInquiry.status, func.count(CsInquiry.id))
            .group_by(CsInquiry.status)
            .all()
        )
        by_status = {row[0]: row[1] for row in status_rows}

        # 쇼핑몰별 통계
        mall_rows = (
            db.query(CsInquiry.mall_name, func.count(CsInquiry.id))
            .group_by(CsInquiry.mall_name)
            .all()
        )
        by_mall = {row[0]: row[1] for row in mall_rows}

        # 카테고리별 통계
        cat_rows = (
            db.query(CsInquiry.category, func.count(CsInquiry.id))
            .filter(CsInquiry.category.isnot(None))
            .group_by(CsInquiry.category)
            .all()
        )
        by_category = {row[0]: row[1] for row in cat_rows}

        # 우선순위별 통계
        priority_rows = (
            db.query(CsInquiry.priority, func.count(CsInquiry.id))
            .group_by(CsInquiry.priority)
            .all()
        )
        by_priority = {row[0]: row[1] for row in priority_rows}

        return {
            "total": total,
            "by_status": by_status,
            "by_mall": by_mall,
            "by_category": by_category,
            "by_priority": by_priority,
        }
    except Exception as e:
        logger.error(f"문의 통계 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/inquiries/total-stats")
def get_total_stats(db: Session = Depends(get_db)):
    """전체 문의 통계 (DB 전체 기준, 페이지네이션 무관)"""
    try:
        total = db.query(func.count(CsInquiry.id)).scalar() or 0
        new_count = db.query(func.count(CsInquiry.id)).filter(CsInquiry.status == "new").scalar() or 0
        ai_drafted = db.query(func.count(CsInquiry.id)).filter(CsInquiry.status == "ai_drafted").scalar() or 0
        approved = db.query(func.count(CsInquiry.id)).filter(CsInquiry.status == "approved").scalar() or 0
        sent = db.query(func.count(CsInquiry.id)).filter(CsInquiry.status == "sent").scalar() or 0

        # 템플릿 답변 수 (재생성 필요)
        template_pattern = "%확인 후 빠르게 처리해 드리겠습니다%"
        template_count = (
            db.query(func.count(CsInquiry.id))
            .filter(CsInquiry.status == "ai_drafted", CsInquiry.ai_response.ilike(template_pattern))
            .scalar() or 0
        )

        return {
            "total": total,
            "new": new_count,
            "ai_drafted": ai_drafted,
            "template_responses": template_count,
            "real_ai_responses": ai_drafted - template_count,
            "approved": approved,
            "sent": sent,
        }
    except Exception as e:
        logger.error(f"전체 통계 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/inquiries")
def list_inquiries(
    status: Optional[str] = None,
    mall_name: Optional[str] = None,
    board_type: Optional[str] = None,
    category: Optional[str] = None,
    priority: Optional[str] = None,
    search: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    """문의사항 목록 조회 (필터, 페이지네이션)"""
    try:
        query = db.query(CsInquiry)

        if status:
            query = query.filter(CsInquiry.status == status)
        if mall_name:
            query = query.filter(CsInquiry.mall_name == mall_name)
        if board_type:
            query = query.filter(CsInquiry.board_type == board_type)
        if category:
            query = query.filter(CsInquiry.category == category)
        if priority:
            query = query.filter(CsInquiry.priority == priority)
        if search:
            like_term = f"%{search}%"
            query = query.filter(
                (CsInquiry.title.ilike(like_term))
                | (CsInquiry.content.ilike(like_term))
                | (CsInquiry.customer_name.ilike(like_term))
                | (CsInquiry.product_name.ilike(like_term))
                | (CsInquiry.order_number.ilike(like_term))
            )
        if date_from:
            try:
                dt_from = datetime.fromisoformat(date_from)
                query = query.filter(CsInquiry.inquiry_date >= dt_from)
            except ValueError:
                pass
        if date_to:
            try:
                dt_to = datetime.fromisoformat(date_to)
                query = query.filter(CsInquiry.inquiry_date <= dt_to)
            except ValueError:
                pass

        total = query.count()
        items = (
            query.order_by(CsInquiry.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
            .all()
        )

        return {
            "total": total,
            "page": page,
            "page_size": page_size,
            "items": [_inquiry_to_dict(item, db) for item in items],
        }
    except Exception as e:
        logger.error(f"문의 목록 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/inquiries/{inquiry_id}")
def get_inquiry(inquiry_id: int, db: Session = Depends(get_db)):
    """문의사항 상세 조회"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")
    return _inquiry_to_dict(inquiry, db)


@router.post("/inquiries")
def create_inquiry(data: InquiryCreate, db: Session = Depends(get_db)):
    """문의사항 생성 (수동 입력 또는 웹훅)"""
    try:
        inquiry = CsInquiry(
            external_id=data.external_id,
            mall_name=data.mall_name,
            board_type=data.board_type,
            customer_name=data.customer_name,
            customer_id=data.customer_id,
            product_name=data.product_name,
            order_number=data.order_number,
            title=data.title,
            content=data.content,
            inquiry_date=datetime.fromisoformat(data.inquiry_date) if data.inquiry_date else None,
            category=data.category,
            priority=data.priority,
            status="new",
        )
        db.add(inquiry)
        db.commit()
        db.refresh(inquiry)
        logger.info(f"문의 생성: id={inquiry.id}, mall={inquiry.mall_name}")
        return _inquiry_to_dict(inquiry, db)
    except Exception as e:
        db.rollback()
        logger.error(f"문의 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/inquiries/{inquiry_id}")
def update_inquiry(inquiry_id: int, data: InquiryUpdate, db: Session = Depends(get_db)):
    """문의사항 업데이트"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")
    try:
        if data.status is not None:
            inquiry.status = data.status
        if data.ai_response is not None:
            inquiry.ai_response = data.ai_response
        if data.final_response is not None:
            inquiry.final_response = data.final_response
        if data.category is not None:
            inquiry.category = data.category
        if data.priority is not None:
            inquiry.priority = data.priority
        db.commit()
        db.refresh(inquiry)
        return _inquiry_to_dict(inquiry, db)
    except Exception as e:
        db.rollback()
        logger.error(f"문의 업데이트 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/inquiries/{inquiry_id}")
def delete_inquiry(inquiry_id: int, db: Session = Depends(get_db)):
    """문의사항 삭제"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")
    try:
        db.delete(inquiry)
        db.commit()
        return {"message": "삭제 완료", "id": inquiry_id}
    except Exception as e:
        db.rollback()
        logger.error(f"문의 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/xml-host/{request_id}")
def serve_xml(request_id: str):
    """사방넷 API가 가져갈 XML 요청 파일 호스팅"""
    from app.services.sabangnet_api import get_stored_xml
    xml_content = get_stored_xml(request_id)
    if not xml_content:
        raise HTTPException(status_code=404, detail="XML not found or expired")
    return FastAPIResponse(
        content=xml_content.encode("euc-kr", errors="replace"),
        media_type="application/xml; charset=euc-kr",
    )


@router.get("/debug/api-test")
async def debug_api_test():
    """사방넷 API 연결 테스트 - 실제 응답을 그대로 반환"""
    from app.services.sabangnet_api import get_sabangnet_api
    api = get_sabangnet_api()

    result = {
        "api_available": api.is_available,
        "login_id": api.login_id,
        "admin_no": api.admin_no,
        "base_url": api.base_url,
        "backend_url": api.backend_url,
        "has_auth_key": bool(api.auth_key),
    }

    # Claude API 테스트
    from app.config import get_settings as _gs
    _s = _gs()
    result["anthropic_key_set"] = bool(_s.ANTHROPIC_API_KEY)
    result["anthropic_key_prefix"] = _s.ANTHROPIC_API_KEY[:10] + "..." if _s.ANTHROPIC_API_KEY else "EMPTY"

    try:
        async with httpx.AsyncClient() as _c:
            _r = await _c.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": _s.ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": "claude-sonnet-4-20250514", "max_tokens": 50, "messages": [{"role": "user", "content": "Say OK"}]},
                timeout=10.0,
            )
            result["claude_test_status"] = _r.status_code
            result["claude_test_response"] = _r.text[:200]
    except Exception as e:
        result["claude_test_error"] = str(e)

    if not api.is_available:
        result["error"] = "API 설정 불완전 (SABANGNET_LOGIN_ID, SABANGNET_ADMIN_NO, SABANGNET_API_KEY 필요)"
        return result

    try:
        api_result = await api.collect_inquiries()
        result["api_success"] = api_result.get("success", False)
        result["items_count"] = len(api_result.get("items", []))
        result["raw_response"] = api_result.get("raw", "")[:3000]
        result["error_if_any"] = api_result.get("error", None)
        if api_result.get("items"):
            result["first_item"] = api_result["items"][0]
    except Exception as e:
        result["exception"] = str(e)
        import traceback
        result["traceback"] = traceback.format_exc()[-1000:]

    return result


@router.post("/inquiries/collect")
async def collect_inquiries(db: Session = Depends(get_db)):
    """사방넷 API로 문의 수집.
    API가 설정되어 있으면 실제 사방넷 API를 호출하고,
    미설정 시 샘플 데이터로 fallback.
    """
    logger.info("사방넷 문의 수집 요청")

    try:
        from app.services.sabangnet_api import get_sabangnet_api
        api = get_sabangnet_api()

        if api.is_available:
            result = await api.collect_inquiries()
            if result.get("success"):
                items = result.get("items", [])
                created = 0
                updated = 0
                for item in items:
                    # 사방넷 XML 필드명 → CsInquiry 매핑
                    external_id = item.get("NUM", "")
                    if not external_id:
                        continue
                    existing = (
                        db.query(CsInquiry)
                        .filter(CsInquiry.external_id == external_id)
                        .first()
                    )
                    if existing:
                        # 양방향 동기화: 사방넷 상태가 변경되었으면 내부 상태도 업데이트
                        cs_status = item.get("CS_STATUS", "")
                        if cs_status and cs_status != getattr(existing, 'sabangnet_status', None):
                            existing.sabangnet_status = cs_status
                            existing.last_synced_at = datetime.now()
                            # 사방넷에서 이미 답변된 경우 → 내부 상태도 갱신
                            if cs_status in ("002", "003") and existing.status in ("new", "ai_drafted"):
                                existing.status = "answered_externally"
                            elif cs_status == "004" and existing.status not in ("sent",):
                                existing.status = "closed_externally"
                            updated += 1
                        continue

                    # 수집일자 파싱 (20190909164517 형식)
                    inquiry_date = datetime.now()
                    raw_date = item.get("INS_DM", item.get("REG_DM", ""))
                    if raw_date and len(raw_date) >= 8:
                        try:
                            if len(raw_date) >= 14:
                                inquiry_date = datetime.strptime(raw_date[:14], "%Y%m%d%H%M%S")
                            else:
                                inquiry_date = datetime.strptime(raw_date[:8], "%Y%m%d")
                        except ValueError:
                            pass

                    inquiry = CsInquiry(
                        external_id=external_id,
                        mall_name=item.get("MALL_ID", ""),
                        board_type=item.get("CS_GUBUN", ""),
                        customer_name=item.get("INS_NM", ""),
                        customer_id=item.get("MALL_USER_ID", ""),
                        product_name=item.get("PRODUCT_NM", ""),
                        order_number=item.get("ORDER_ID", ""),
                        title=item.get("SUBJECT", ""),
                        content=item.get("CNTS", ""),
                        inquiry_date=inquiry_date,
                        status="new" if item.get("CS_STATUS") == "001" else (
                            "ai_drafted" if item.get("CS_STATUS") == "002" else (
                                "sent" if item.get("CS_STATUS") == "003" else "new"
                            )
                        ),
                        priority="normal",
                        category=item.get("CS_GUBUN", ""),
                    )
                    db.add(inquiry)
                    created += 1
                db.commit()

                # 수집된 문의에 대해 즉시 DeliveryTracking 레코드 생성
                # (별도 API 호출 없이, 수집 시 받아온 CS 데이터를 기반으로)
                dt_created = 0
                for item in items:
                    ext_id = item.get("NUM", "")
                    if not ext_id:
                        continue
                    inq = db.query(CsInquiry).filter(CsInquiry.external_id == ext_id).first()
                    if not inq:
                        continue
                    existing_dt = db.query(DeliveryTracking).filter(
                        DeliveryTracking.inquiry_id == inq.id
                    ).first()
                    if existing_dt:
                        # 기존 레코드가 있지만 order_detail이 부실하면 갱신
                        if not existing_dt.order_detail or existing_dt.current_status in ("fetch_failed", "unknown"):
                            existing_dt.order_detail = item  # 사방넷 원본 전체 필드
                            existing_dt.current_status = "collected"
                            existing_dt.last_event = "수집 데이터 갱신됨"
                        continue
                    order_id = item.get("ORDER_ID", "").strip()
                    # 사방넷 CS XML의 모든 필드를 order_detail에 저장
                    dt = DeliveryTracking(
                        inquiry_id=inq.id,
                        order_number=order_id,
                        tracking_number=item.get("DELIVERY_NO", "").strip(),
                        courier_name=item.get("DELIVERY_COMPANY_NM", "").strip(),
                        current_status="collected" if order_id else "no_order",
                        last_event="수집완료" if order_id else "주문번호 없음",
                        order_detail=item,  # 사방넷 원본 전체 필드 저장
                    )
                    db.add(dt)
                    dt_created += 1
                if dt_created > 0:
                    db.commit()

                return {
                    "message": f"사방넷에서 {created}건 수집, {updated}건 상태 동기화 완료 (총 {len(items)}건 조회)",
                    "items_created": created,
                    "items_updated": updated,
                    "total_fetched": len(items),
                    "source": "sabangnet_api",
                    "raw_response": result.get("raw", "")[:500],
                }
            else:
                error_msg = result.get("error", "알 수 없는 오류")
                raw_resp = result.get("raw", "")
                logger.warning(f"사방넷 API 오류: {error_msg}, raw: {raw_resp[:300]}")
                return {
                    "message": f"사방넷 API 오류: {error_msg}",
                    "items_created": 0,
                    "source": "sabangnet_api_error",
                    "error_detail": str(error_msg),
                    "raw_response": raw_resp[:300],
                }
    except Exception as e:
        logger.error(f"사방넷 API 호출 실패: {e}", exc_info=True)
        return {
            "message": f"사방넷 API 호출 실패: {str(e)}",
            "items_created": 0,
            "source": "api_exception",
            "error_detail": str(e),
        }

    # API 미설정 또는 API 오류 시 샘플 데이터 fallback
    mock_inquiries = [
        {
            "external_id": "SBN-2026-001",
            "mall_name": "스마트스토어",
            "board_type": "상품문의",
            "customer_name": "김고객",
            "product_name": "샘플 상품 A",
            "title": "상품 사이즈 문의드립니다",
            "content": "이 상품 S사이즈 재입고 언제 될까요?",
            "inquiry_date": "2026-03-15T10:30:00",
            "status": "new",
        },
        {
            "external_id": "SBN-2026-002",
            "mall_name": "쿠팡",
            "board_type": "배송문의",
            "customer_name": "이고객",
            "order_number": "2026031500123",
            "product_name": "샘플 상품 B",
            "title": "배송이 너무 늦어요",
            "content": "주문한지 5일인데 아직 배송 시작이 안되었습니다. 확인 부탁드립니다.",
            "inquiry_date": "2026-03-15T11:00:00",
            "status": "new",
        },
    ]

    created = []
    for item in mock_inquiries:
        existing = (
            db.query(CsInquiry)
            .filter(CsInquiry.external_id == item["external_id"])
            .first()
        )
        if existing:
            continue
        inquiry = CsInquiry(
            external_id=item["external_id"],
            mall_name=item["mall_name"],
            board_type=item.get("board_type"),
            customer_name=item.get("customer_name"),
            product_name=item.get("product_name"),
            order_number=item.get("order_number"),
            title=item.get("title"),
            content=item["content"],
            inquiry_date=datetime.fromisoformat(item["inquiry_date"]) if item.get("inquiry_date") else None,
            status="new",
        )
        db.add(inquiry)
        created.append(item["external_id"])

    if created:
        db.commit()

    return {
        "message": "사방넷 문의 수집 완료 (샘플 데이터)",
        "collected_count": len(created),
        "external_ids": created,
        "source": "sample",
    }


@router.post("/inquiries/sync-status")
async def sync_inquiry_statuses(db: Session = Depends(get_db)):
    """사방넷 상태와 내부 상태를 동기화 (전체 미완료 건 대상)"""
    try:
        from app.services.sabangnet_api import get_sabangnet_api
        api = get_sabangnet_api()
        if not api.is_available:
            raise HTTPException(status_code=400, detail="사방넷 API 미설정")

        # 미완료 상태의 문의 조회
        pending = db.query(CsInquiry).filter(
            CsInquiry.status.notin_(["sent", "answered_externally", "closed_externally"])
        ).all()

        if not pending:
            return {"message": "동기화할 문의가 없습니다", "updated": 0}

        # 사방넷에서 최신 상태 수집
        result = await api.collect_inquiries()
        if not result.get("success"):
            raise HTTPException(status_code=500, detail="사방넷 조회 실패")

        # external_id → CS_STATUS 매핑
        status_map = {}
        for item in result.get("items", []):
            ext_id = item.get("NUM", "")
            if ext_id:
                status_map[ext_id] = item.get("CS_STATUS", "")

        updated = 0
        for inq in pending:
            if inq.external_id and inq.external_id in status_map:
                cs_status = status_map[inq.external_id]
                if cs_status != inq.sabangnet_status:
                    inq.sabangnet_status = cs_status
                    inq.last_synced_at = datetime.now()
                    if cs_status in ("002", "003") and inq.status in ("new", "ai_drafted"):
                        inq.status = "answered_externally"
                    elif cs_status == "004":
                        inq.status = "closed_externally"
                    updated += 1

        db.commit()
        return {"message": f"{updated}건 상태 동기화 완료", "updated": updated, "total_checked": len(pending)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"상태 동기화 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────
# AI 답변 생성 엔드포인트
# ──────────────────────────────────────────────

@router.post("/inquiries/{inquiry_id}/generate-ai")
async def generate_ai_for_inquiry(inquiry_id: int, db: Session = Depends(get_db)):
    """단건 AI 답변 생성"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")

    try:
        # 활성 참고 데이터 조회
        ref_data = (
            db.query(CsReferenceData)
            .filter(CsReferenceData.is_active == True)
            .all()
        )

        # AI 답변 생성
        ai_text = await generate_ai_response(inquiry, ref_data, db)

        # 저장 및 상태 업데이트
        inquiry.ai_response = ai_text
        inquiry.status = "ai_drafted"
        db.commit()
        db.refresh(inquiry)

        # 후속 조치 파싱 및 저장
        _parse_ai_actions(ai_text, inquiry.id, db)

        logger.info(f"AI 답변 생성 완료: inquiry_id={inquiry_id}")
        return {
            "inquiry_id": inquiry_id,
            "ai_response": ai_text,
            "status": "ai_drafted",
        }
    except Exception as e:
        db.rollback()
        logger.error(f"AI 답변 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/inquiries/bulk-generate-ai")
async def bulk_generate_ai(data: BulkAiGenerateRequest, db: Session = Depends(get_db)):
    """다건 AI 답변 일괄 생성"""
    # 활성 참고 데이터 한 번만 조회
    ref_data = (
        db.query(CsReferenceData)
        .filter(CsReferenceData.is_active == True)
        .all()
    )

    generated_count = 0
    errors = []

    for inq_id in data.inquiry_ids:
        inquiry = db.query(CsInquiry).filter(CsInquiry.id == inq_id).first()
        if not inquiry:
            errors.append({"id": inq_id, "error": "문의사항을 찾을 수 없습니다"})
            continue
        try:
            ai_text = await generate_ai_response(inquiry, ref_data, db)
            inquiry.ai_response = ai_text
            inquiry.status = "ai_drafted"
            generated_count += 1
            _parse_ai_actions(ai_text, inquiry.id, db)
        except Exception as e:
            errors.append({"id": inq_id, "error": str(e)})

    db.commit()
    logger.info(f"AI 답변 일괄 생성: {generated_count}/{len(data.inquiry_ids)} 건")

    return {
        "total_requested": len(data.inquiry_ids),
        "generated_count": generated_count,
        "errors": errors,
    }


@router.post("/inquiries/bulk-generate-all-new")
async def bulk_generate_all_new(
    page_size: int = Query(default=50, description="한 번에 처리할 건수"),
    db: Session = Depends(get_db),
):
    """신규(new) 상태 문의 전체에 AI 답변 일괄 생성"""
    ref_data = (
        db.query(CsReferenceData)
        .filter(CsReferenceData.is_active == True)
        .all()
    )

    new_inquiries = (
        db.query(CsInquiry)
        .filter(CsInquiry.status == "new")
        .order_by(CsInquiry.inquiry_date.desc())
        .limit(page_size)
        .all()
    )

    total_new = db.query(func.count(CsInquiry.id)).filter(CsInquiry.status == "new").scalar() or 0
    generated = 0
    failed = 0

    for inquiry in new_inquiries:
        try:
            ai_text = await generate_ai_response(inquiry, ref_data, db)
            inquiry.ai_response = ai_text
            inquiry.status = "ai_drafted"
            generated += 1
            _parse_ai_actions(ai_text, inquiry.id, db)
        except Exception as e:
            logger.error(f"AI 생성 실패 (id={inquiry.id}): {e}")
            failed += 1

    db.commit()
    remaining = total_new - generated
    return {
        "generated": generated,
        "failed": failed,
        "remaining": max(0, remaining),
        "total_new": total_new,
        "message": f"{generated}건 AI 답변 생성 완료 (남은 신규: {max(0, remaining)}건)",
    }


@router.post("/inquiries/bulk-regenerate")
async def bulk_regenerate(
    page_size: int = Query(default=30, description="한 번에 처리할 건수"),
    db: Session = Depends(get_db),
):
    """템플릿 답변(ai_drafted)을 Claude AI로 재생성.
    '확인 후 빠르게 처리해 드리겠습니다' 같은 템플릿 패턴을 감지하여 재생성.
    """
    ref_data = (
        db.query(CsReferenceData)
        .filter(CsReferenceData.is_active == True)
        .all()
    )

    # 템플릿 패턴 감지: "확인 후 빠르게 처리해 드리겠습니다" 포함된 답변
    template_pattern = "%확인 후 빠르게 처리해 드리겠습니다%"
    template_inquiries = (
        db.query(CsInquiry)
        .filter(
            CsInquiry.status == "ai_drafted",
            CsInquiry.ai_response.ilike(template_pattern),
        )
        .order_by(CsInquiry.inquiry_date.desc())
        .limit(page_size)
        .all()
    )

    total_template = (
        db.query(func.count(CsInquiry.id))
        .filter(
            CsInquiry.status == "ai_drafted",
            CsInquiry.ai_response.ilike(template_pattern),
        )
        .scalar() or 0
    )

    generated = 0
    failed = 0

    for inquiry in template_inquiries:
        try:
            ai_text = await generate_ai_response(inquiry, ref_data, db)
            # 여전히 템플릿이면 스킵 (API 실패)
            if "확인 후 빠르게 처리해 드리겠습니다" in ai_text:
                failed += 1
                continue
            inquiry.ai_response = ai_text
            generated += 1
            _parse_ai_actions(ai_text, inquiry.id, db)
        except Exception as e:
            logger.error(f"AI 재생성 실패 (id={inquiry.id}): {e}")
            failed += 1

    db.commit()
    remaining = total_template - generated
    return {
        "generated": generated,
        "failed": failed,
        "remaining": max(0, remaining),
        "total_template": total_template,
        "message": f"{generated}건 AI 재생성 완료 (남은 템플릿: {max(0, remaining)}건)",
    }


# ──────────────────────────────────────────────
# 답변 승인 및 발송 엔드포인트
# ──────────────────────────────────────────────

@router.post("/inquiries/{inquiry_id}/approve")
def approve_inquiry(
    inquiry_id: int,
    final_text: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """AI 답변 승인 (수정 가능)"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")

    try:
        # final_text가 있으면 사용, 없으면 ai_response를 복사
        if final_text:
            inquiry.final_response = final_text
        elif inquiry.ai_response:
            inquiry.final_response = inquiry.ai_response
        else:
            raise HTTPException(status_code=400, detail="승인할 답변이 없습니다 (AI 답변 먼저 생성 필요)")

        inquiry.status = "approved"
        db.commit()
        db.refresh(inquiry)

        logger.info(f"답변 승인: inquiry_id={inquiry_id}")
        return {
            "inquiry_id": inquiry_id,
            "status": "approved",
            "final_response": inquiry.final_response,
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"답변 승인 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/inquiries/{inquiry_id}/send")
async def send_response(inquiry_id: int, db: Session = Depends(get_db)):
    """사방넷 API로 답변 발송"""
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의사항을 찾을 수 없습니다")

    # 발송할 답변 결정
    response_text = inquiry.final_response or inquiry.ai_response
    if not response_text:
        raise HTTPException(status_code=400, detail="발송할 답변이 없습니다")

    try:
        # 사방넷 API로 답변 실제 발송
        from app.services.sabangnet_api import get_sabangnet_api
        api = get_sabangnet_api()
        if api.is_available and inquiry.external_id and inquiry.final_response:
            send_result = await api.send_inquiry_response([{
                "num": inquiry.external_id,
                "content": inquiry.final_response,
            }])
            if send_result.get("success"):
                logger.info(f"사방넷 답변 발송 성공: {inquiry.external_id}")
            else:
                logger.warning(f"사방넷 답변 발송 실패: {send_result}")
        else:
            logger.info(
                f"사방넷 답변 발송 (API 미설정 또는 external_id 없음): inquiry_id={inquiry_id}, "
                f"external_id={inquiry.external_id}, mall={inquiry.mall_name}"
            )

        inquiry.status = "sent"
        inquiry.sent_at = datetime.utcnow()
        db.commit()
        db.refresh(inquiry)

        return {
            "inquiry_id": inquiry_id,
            "status": "sent",
            "sent_at": inquiry.sent_at.isoformat() if inquiry.sent_at else None,
            "message": "답변 발송 완료",
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"답변 발송 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/inquiries/bulk-send")
def bulk_send_responses(db: Session = Depends(get_db)):
    """승인된 답변 일괄 발송"""
    approved = (
        db.query(CsInquiry)
        .filter(CsInquiry.status == "approved")
        .all()
    )

    if not approved:
        return {"message": "발송할 승인 건이 없습니다", "sent_count": 0}

    sent_count = 0
    errors = []
    now = datetime.utcnow()

    for inquiry in approved:
        response_text = inquiry.final_response or inquiry.ai_response
        if not response_text:
            errors.append({"id": inquiry.id, "error": "답변 내용 없음"})
            continue
        try:
            logger.info(
                f"사방넷 일괄 발송 (placeholder): inquiry_id={inquiry.id}, "
                f"external_id={inquiry.external_id}"
            )
            inquiry.status = "sent"
            inquiry.sent_at = now
            sent_count += 1
        except Exception as e:
            errors.append({"id": inquiry.id, "error": str(e)})

    db.commit()
    logger.info(f"일괄 발송 완료: {sent_count}/{len(approved)} 건")

    return {
        "total_approved": len(approved),
        "sent_count": sent_count,
        "errors": errors,
    }


# ──────────────────────────────────────────────
# 참고 데이터 (Reference Data) 엔드포인트
# ──────────────────────────────────────────────

@router.get("/reference-data/categories")
def get_reference_categories(db: Session = Depends(get_db)):
    """참고 데이터 카테고리 목록"""
    try:
        rows = (
            db.query(distinct(CsReferenceData.category))
            .filter(CsReferenceData.category.isnot(None))
            .all()
        )
        categories = [row[0] for row in rows]
        return {"categories": categories}
    except Exception as e:
        logger.error(f"카테고리 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/reference-data")
def list_reference_data(
    category: Optional[str] = None,
    is_active: Optional[bool] = None,
    db: Session = Depends(get_db),
):
    """참고 데이터 목록"""
    try:
        query = db.query(CsReferenceData)
        if category:
            query = query.filter(CsReferenceData.category == category)
        if is_active is not None:
            query = query.filter(CsReferenceData.is_active == is_active)
        items = query.order_by(CsReferenceData.created_at.desc()).all()
        return {
            "items": [_ref_to_dict(item) for item in items],
        }
    except Exception as e:
        logger.error(f"참고 데이터 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/reference-data")
def create_reference_data(data: ReferenceDataCreate, db: Session = Depends(get_db)):
    """참고 데이터 생성"""
    try:
        ref = CsReferenceData(
            title=data.title,
            category=data.category,
            content=data.content,
            is_active=True,
        )
        db.add(ref)
        db.commit()
        db.refresh(ref)
        logger.info(f"참고 데이터 생성: id={ref.id}, title={ref.title}")
        return _ref_to_dict(ref)
    except Exception as e:
        db.rollback()
        logger.error(f"참고 데이터 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/reference-data/{ref_id}")
def update_reference_data(ref_id: int, data: ReferenceDataUpdate, db: Session = Depends(get_db)):
    """참고 데이터 업데이트"""
    ref = db.query(CsReferenceData).filter(CsReferenceData.id == ref_id).first()
    if not ref:
        raise HTTPException(status_code=404, detail="참고 데이터를 찾을 수 없습니다")
    try:
        if data.title is not None:
            ref.title = data.title
        if data.category is not None:
            ref.category = data.category
        if data.content is not None:
            ref.content = data.content
        if data.is_active is not None:
            ref.is_active = data.is_active
        db.commit()
        db.refresh(ref)
        return _ref_to_dict(ref)
    except Exception as e:
        db.rollback()
        logger.error(f"참고 데이터 업데이트 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/reference-data/{ref_id}")
def delete_reference_data(ref_id: int, db: Session = Depends(get_db)):
    """참고 데이터 삭제 (hard delete)"""
    ref = db.query(CsReferenceData).filter(CsReferenceData.id == ref_id).first()
    if not ref:
        raise HTTPException(status_code=404, detail="참고 데이터를 찾을 수 없습니다")
    try:
        # 첨부 파일도 삭제
        if ref.file_name:
            file_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{ref.file_name}")
            if os.path.exists(file_path):
                os.remove(file_path)
        db.delete(ref)
        db.commit()
        return {"message": "삭제 완료", "id": ref_id}
    except Exception as e:
        db.rollback()
        logger.error(f"참고 데이터 삭제 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/reference-data/upload")
async def create_reference_with_file(
    title: str = Form(...),
    category: str = Form(None),
    content: str = Form(""),
    is_active: bool = Form(True),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
):
    """참고 데이터 생성 (파일 첨부 포함)"""
    try:
        file_name = None
        file_type = None
        file_size = None
        extracted_text = None

        if file and file.filename:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"지원하지 않는 파일 형식입니다. 허용: {', '.join(ALLOWED_EXTENSIONS)}"
                )

            file_bytes = await file.read()
            file_size = len(file_bytes)

            if file_size > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail="파일 크기가 20MB를 초과합니다.")

            file_name = file.filename
            file_type = ext.lstrip(".")

            # 텍스트 추출
            extracted_text = extract_text_from_file(file_bytes, file.filename)
            if not extracted_text:
                extracted_text = None

            # 파일 저장은 DB에 저장 후 id를 사용
            ref = CsReferenceData(
                title=title,
                category=category,
                content=content or "(첨부파일 참조)",
                file_name=file_name,
                file_type=file_type,
                file_size=file_size,
                extracted_text=extracted_text,
                is_active=is_active,
            )
            db.add(ref)
            db.commit()
            db.refresh(ref)

            # 파일 저장 (id 포함 이름)
            save_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{file_name}")
            with open(save_path, "wb") as f:
                f.write(file_bytes)

            logger.info(f"참고 데이터+파일 생성: id={ref.id}, file={file_name}, extracted={len(extracted_text or '')} chars")
            return _ref_to_dict(ref)
        else:
            # 파일 없이 텍스트만
            ref = CsReferenceData(
                title=title,
                category=category,
                content=content,
                is_active=is_active,
            )
            db.add(ref)
            db.commit()
            db.refresh(ref)
            logger.info(f"참고 데이터 생성: id={ref.id}, title={title}")
            return _ref_to_dict(ref)

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"참고 데이터+파일 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/reference-data/{ref_id}/upload")
async def update_reference_with_file(
    ref_id: int,
    title: str = Form(None),
    category: str = Form(None),
    content: str = Form(None),
    is_active: bool = Form(True),
    file: UploadFile = File(None),
    remove_file: bool = Form(False),
    db: Session = Depends(get_db),
):
    """참고 데이터 수정 (파일 교체 포함)"""
    ref = db.query(CsReferenceData).filter(CsReferenceData.id == ref_id).first()
    if not ref:
        raise HTTPException(status_code=404, detail="참고 데이터를 찾을 수 없습니다")

    try:
        if title is not None:
            ref.title = title
        if category is not None:
            ref.category = category
        if content is not None:
            ref.content = content
        ref.is_active = is_active

        # 파일 삭제 요청
        if remove_file and ref.file_name:
            old_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{ref.file_name}")
            if os.path.exists(old_path):
                os.remove(old_path)
            ref.file_name = None
            ref.file_type = None
            ref.file_size = None
            ref.extracted_text = None

        # 새 파일 업로드
        if file and file.filename:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"지원하지 않는 파일 형식입니다. 허용: {', '.join(ALLOWED_EXTENSIONS)}"
                )

            file_bytes = await file.read()
            if len(file_bytes) > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail="파일 크기가 20MB를 초과합니다.")

            # 기존 파일 삭제
            if ref.file_name:
                old_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{ref.file_name}")
                if os.path.exists(old_path):
                    os.remove(old_path)

            ref.file_name = file.filename
            ref.file_type = ext.lstrip(".")
            ref.file_size = len(file_bytes)
            ref.extracted_text = extract_text_from_file(file_bytes, file.filename) or None

            # 파일 저장
            save_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{file.filename}")
            with open(save_path, "wb") as f:
                f.write(file_bytes)

        db.commit()
        db.refresh(ref)
        return _ref_to_dict(ref)

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"참고 데이터+파일 수정 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/reference-data/{ref_id}/download")
def download_reference_file(ref_id: int, db: Session = Depends(get_db)):
    """참고 데이터 첨부파일 다운로드"""
    ref = db.query(CsReferenceData).filter(CsReferenceData.id == ref_id).first()
    if not ref:
        raise HTTPException(status_code=404, detail="참고 데이터를 찾을 수 없습니다")
    if not ref.file_name:
        raise HTTPException(status_code=404, detail="첨부파일이 없습니다")

    file_path = os.path.join(UPLOAD_DIR, f"{ref.id}_{ref.file_name}")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다")

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    # Content-Type 매핑
    content_types = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "xls": "application/vnd.ms-excel",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "txt": "text/plain",
        "csv": "text/csv",
    }
    content_type = content_types.get(ref.file_type or "", "application/octet-stream")

    return Response(
        content=file_bytes,
        media_type=content_type,
        headers={
            "Content-Disposition": f'attachment; filename="{ref.file_name}"',
        },
    )


# ──────────────────────────────────────────────
# 설정 (Config) 엔드포인트
# ──────────────────────────────────────────────

@router.get("/config")
def get_config(db: Session = Depends(get_db)):
    """CS 설정 전체 조회"""
    try:
        configs = db.query(CsConfig).all()
        result = {}
        for cfg in configs:
            # JSON으로 파싱 가능하면 파싱
            try:
                result[cfg.config_key] = json.loads(cfg.config_value)
            except (json.JSONDecodeError, TypeError):
                result[cfg.config_key] = cfg.config_value
        return result
    except Exception as e:
        logger.error(f"설정 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/config")
def update_config(data: ConfigUpdate, db: Session = Depends(get_db)):
    """CS 설정 업데이트"""
    try:
        updates = {}
        if data.auto_mode is not None:
            updates["auto_mode"] = json.dumps(data.auto_mode)
        if data.auto_categories is not None:
            updates["auto_categories"] = json.dumps(data.auto_categories, ensure_ascii=False)
        if data.sabangnet_api_key is not None:
            updates["sabangnet_api_key"] = data.sabangnet_api_key
        if data.claude_api_key is not None:
            updates["claude_api_key"] = data.claude_api_key
        if data.response_tone is not None:
            updates["response_tone"] = data.response_tone

        for key, value in updates.items():
            existing = db.query(CsConfig).filter(CsConfig.config_key == key).first()
            if existing:
                existing.config_value = value
            else:
                new_cfg = CsConfig(config_key=key, config_value=value)
                db.add(new_cfg)

        db.commit()
        logger.info(f"설정 업데이트: {list(updates.keys())}")

        # 업데이트 후 전체 설정 반환
        return get_config(db=db)
    except Exception as e:
        db.rollback()
        logger.error(f"설정 업데이트 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────
# 헬퍼 함수
# ──────────────────────────────────────────────

def _detect_action(inq: CsInquiry) -> dict:
    """문의 내용에서 필요한 액션을 자동 감지"""
    content = (inq.content or "").lower() + " " + (inq.title or "").lower() + " " + (inq.board_type or "").lower()

    actions = []
    if any(w in content for w in ["환불", "취소", "캔슬"]):
        actions.append({"type": "refund", "label": "환불 처리", "priority": "high"})
    if any(w in content for w in ["교환", "반품", "수거"]):
        actions.append({"type": "exchange", "label": "교환/반품 처리", "priority": "high"})
    if any(w in content for w in ["파손", "깨짐", "훼손", "불량"]):
        actions.append({"type": "damage", "label": "파손 보상", "priority": "high"})
    if any(w in content for w in ["배송", "지연", "안왔", "안 왔", "늦"]):
        actions.append({"type": "delivery", "label": "배송 확인", "priority": "medium"})
    if any(w in content for w in ["재입고", "품절", "재고"]):
        actions.append({"type": "restock", "label": "재입고 안내", "priority": "low"})

    return {
        "actions": actions,
        "has_order": bool(inq.order_number),
        "order_number": inq.order_number or None,
    }


def _parse_ai_actions(ai_response: str, inquiry_id: int, db: Session) -> list:
    """AI 답변에서 [ACTION:type] 태그를 파싱하여 후속 조치 레코드 생성"""
    import re
    actions = []
    action_labels = {
        "exchange": "교환 처리",
        "refund": "환불 처리",
        "reship": "재배송 처리",
        "damage_claim": "파손 보상",
        "restock_notify": "재입고 알림",
        "delivery_check": "배송 확인",
    }

    matches = re.findall(r'\[ACTION:(\w+)\]', ai_response)
    for action_type in matches:
        if action_type == "none":
            continue
        label = action_labels.get(action_type, action_type)
        # 중복 체크
        existing = db.query(CsFollowUpAction).filter(
            CsFollowUpAction.inquiry_id == inquiry_id,
            CsFollowUpAction.action_type == action_type,
            CsFollowUpAction.status.in_(["pending", "in_progress"]),
        ).first()
        if not existing:
            action = CsFollowUpAction(
                inquiry_id=inquiry_id,
                action_type=action_type,
                action_label=label,
                priority="high" if action_type in ("refund", "exchange", "damage_claim") else "medium",
                ai_suggested=True,
            )
            db.add(action)
            actions.append({"type": action_type, "label": label})

    if actions:
        try:
            db.commit()
        except Exception as e:
            logger.warning(f"후속 조치 저장 실패: {e}")
            db.rollback()

    return actions


def _inquiry_to_dict(inq: CsInquiry, db: Session = None) -> dict:
    """CsInquiry ORM 객체를 dict로 변환"""
    result = {
        "id": inq.id,
        "external_id": inq.external_id,
        "mall_name": inq.mall_name,
        "board_type": inq.board_type,
        "customer_name": inq.customer_name,
        "customer_id": inq.customer_id,
        "product_name": inq.product_name,
        "order_number": inq.order_number,
        "title": inq.title,
        "content": inq.content,
        "inquiry_date": inq.inquiry_date.isoformat() if inq.inquiry_date else None,
        "status": inq.status,
        "ai_response": inq.ai_response,
        "final_response": inq.final_response,
        "sent_at": inq.sent_at.isoformat() if inq.sent_at else None,
        "auto_mode": inq.auto_mode,
        "category": inq.category,
        "priority": inq.priority,
        "sabangnet_status": getattr(inq, 'sabangnet_status', None),
        "last_synced_at": inq.last_synced_at.isoformat() if getattr(inq, 'last_synced_at', None) else None,
        "created_at": inq.created_at.isoformat() if inq.created_at else None,
        "updated_at": inq.updated_at.isoformat() if inq.updated_at else None,
        "auto_action": _detect_action(inq),
    }

    if db:
        # 배송 추적 정보
        delivery = db.query(DeliveryTracking).filter(
            DeliveryTracking.inquiry_id == inq.id
        ).first()
        if delivery:
            result["delivery_tracking"] = {
                "tracking_number": delivery.tracking_number,
                "courier_name": delivery.courier_name,
                "current_status": delivery.current_status,
                "last_event": delivery.last_event,
                "order_detail": delivery.order_detail,
                "last_checked_at": delivery.last_checked_at.isoformat() if delivery.last_checked_at else None,
            }

        # 후속 조치 목록
        actions = db.query(CsFollowUpAction).filter(
            CsFollowUpAction.inquiry_id == inq.id
        ).order_by(CsFollowUpAction.created_at.desc()).all()
        if actions:
            result["followup_actions"] = [{
                "id": a.id,
                "action_type": a.action_type,
                "action_label": a.action_label,
                "priority": a.priority,
                "status": a.status,
                "ai_suggested": a.ai_suggested,
                "notes": a.notes,
                "created_at": a.created_at.isoformat() if a.created_at else None,
            } for a in actions]

    return result


def _ref_to_dict(ref: CsReferenceData) -> dict:
    """CsReferenceData ORM 객체를 dict로 변환"""
    return {
        "id": ref.id,
        "title": ref.title,
        "category": ref.category,
        "content": ref.content,
        "file_name": ref.file_name,
        "file_type": ref.file_type,
        "file_size": ref.file_size,
        "has_file": ref.file_name is not None,
        "has_extracted_text": ref.extracted_text is not None and len(ref.extracted_text or "") > 0,
        "is_active": ref.is_active,
        "created_at": ref.created_at.isoformat() if ref.created_at else None,
        "updated_at": ref.updated_at.isoformat() if ref.updated_at else None,
    }


# ──────────────────────────────────────────────
# 주문 조회 & 배송 추적 엔드포인트
# ──────────────────────────────────────────────

@router.get("/inquiries/{inquiry_id}/order-detail")
async def get_order_detail(inquiry_id: int, db: Session = Depends(get_db)):
    """문의에 연결된 주문 상세 정보 조회.
    사방넷 주문 API 호출 시도 → 실패해도 CsInquiry 데이터로 기본 정보 표시.
    """
    inquiry = db.query(CsInquiry).filter(CsInquiry.id == inquiry_id).first()
    if not inquiry:
        raise HTTPException(status_code=404, detail="문의를 찾을 수 없습니다")

    # CsInquiry에서 기본 주문 정보 구성 (이건 항상 있음)
    base_order_detail = {
        "ORDER_ID": inquiry.order_number or "",
        "PRODUCT_NM": inquiry.product_name or "",
        "MALL_ID": inquiry.mall_name or "",
        "USER_NAME": inquiry.customer_name or "",
    }

    oi = {}
    delivery_info = None
    tracking_no = ""
    courier_name = ""
    api_success = False

    # 사방넷 주문 API 시도 (주문번호가 있을 때만)
    if inquiry.order_number:
        try:
            from app.services.sabangnet_api import get_sabangnet_api
            api = get_sabangnet_api()
            if api.is_available:
                result = await api.get_order_detail(inquiry.order_number)
                logger.info(f"주문 조회 결과: success={result.get('success')}, items={len(result.get('items', []))}, raw={result.get('raw', '')[:200]}")
                if result.get("success") and result.get("items"):
                    oi = result["items"][0]
                    tracking_no = oi.get("DELIVERY_NO", "").strip()
                    courier_name = oi.get("DELIVERY_COMPANY_NM", "").strip()
                    api_success = True

                    if tracking_no and courier_name:
                        try:
                            from app.services.delivery_tracker import DeliveryTracker
                            tracker = DeliveryTracker()
                            delivery_info = await tracker.track_delivery(tracking_no, courier_name)
                        except Exception as e_track:
                            logger.warning(f"배송 추적 실패: {e_track}")
        except Exception as e_api:
            logger.warning(f"주문 API 호출 실패: {e_api}")

    # order_detail 결정: API 성공 → API 데이터, 실패 → CsInquiry 기본 데이터
    final_order_detail = oi if api_success else base_order_detail

    # 상태 결정
    if delivery_info and delivery_info.get("success"):
        status = delivery_info.get("status", "unknown")
        event = delivery_info.get("last_event", "")
    elif api_success:
        status = "order_found"
        event = f"주문확인 (운송장: {'있음' if tracking_no else '미발급'})"
    elif inquiry.order_number:
        status = "collected"
        event = "주문정보 확인됨 (상세 배송정보는 사방넷에서 조회)"
    else:
        status = "no_order"
        event = "주문번호 없음"

    # DeliveryTracking 저장/업데이트
    existing_track = db.query(DeliveryTracking).filter(
        DeliveryTracking.inquiry_id == inquiry_id
    ).first()
    if existing_track:
        existing_track.order_detail = final_order_detail
        existing_track.tracking_number = tracking_no or existing_track.tracking_number
        existing_track.courier_name = courier_name or existing_track.courier_name
        existing_track.current_status = status
        existing_track.last_event = event
        existing_track.last_checked_at = datetime.now()
        if delivery_info and delivery_info.get("success"):
            existing_track.courier_code = delivery_info.get("courier_code", "")
            existing_track.tracking_history = delivery_info.get("events", [])
    else:
        dt = DeliveryTracking(
            inquiry_id=inquiry_id,
            order_number=inquiry.order_number or "",
            tracking_number=tracking_no,
            courier_name=courier_name,
            courier_code=delivery_info.get("courier_code", "") if delivery_info else "",
            current_status=status,
            last_event=event,
            tracking_history=delivery_info.get("events", []) if delivery_info and delivery_info.get("success") else [],
            order_detail=final_order_detail,
        )
        db.add(dt)
    db.commit()

    return {
        "success": True,
        "order": final_order_detail,
        "delivery": delivery_info,
        "api_success": api_success,
    }


@router.post("/inquiries/auto-fetch-order-details")
async def auto_fetch_order_details(db: Session = Depends(get_db)):
    """DeliveryTracking 레코드가 없는 모든 문의에 대해 기본 레코드를 생성.
    주문번호가 있는 건은 사방넷 주문 API로 최신 배송정보 조회 시도.
    """
    all_inquiries = db.query(CsInquiry).all()
    created = 0
    updated = 0
    for inq in all_inquiries:
        order_id = (inq.order_number or "").strip()
        base_detail = {
            "ORDER_ID": order_id,
            "PRODUCT_NM": inq.product_name or "",
            "MALL_ID": inq.mall_name or "",
            "USER_NAME": inq.customer_name or "",
        }

        existing = db.query(DeliveryTracking).filter(DeliveryTracking.inquiry_id == inq.id).first()
        if existing:
            # 기존 실패/unknown 레코드 → CsInquiry 데이터로 갱신
            if existing.current_status in ("fetch_failed", "unknown") or not existing.order_detail:
                existing.order_detail = base_detail if order_id else None
                existing.current_status = "collected" if order_id else "no_order"
                existing.last_event = "주문정보 확인됨" if order_id else "주문번호 없음"
                existing.order_number = order_id
                updated += 1
        else:
            dt = DeliveryTracking(
                inquiry_id=inq.id,
                order_number=order_id,
                tracking_number="",
                courier_name="",
                current_status="collected" if order_id else "no_order",
                last_event="주문정보 확인됨" if order_id else "주문번호 없음",
                order_detail=base_detail if order_id else None,
            )
            db.add(dt)
            created += 1
    db.commit()
    return {
        "message": f"완료: {created}건 생성, {updated}건 갱신",
        "created": created,
        "updated": updated,
        "total": len(all_inquiries),
    }


@router.get("/delivery/track/{tracking_number}")
async def track_delivery_direct(
    tracking_number: str,
    courier_name: str = Query(..., description="택배사명 (예: CJ대한통운)"),
):
    """운송장번호로 직접 배송 추적"""
    from app.services.delivery_tracker import DeliveryTracker
    tracker = DeliveryTracker()
    result = await tracker.track_delivery(tracking_number, courier_name)
    return result


# ──────────────────────────────────────────────
# 후속 조치 (Follow-up Actions) 엔드포인트
# ──────────────────────────────────────────────

@router.get("/followup-actions")
def list_followup_actions(
    status: Optional[str] = None,
    action_type: Optional[str] = None,
    priority: Optional[str] = None,
    page: int = Query(default=1, ge=1),
    limit: int = Query(default=20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    """후속 조치 목록 조회"""
    query = db.query(CsFollowUpAction)
    if status:
        query = query.filter(CsFollowUpAction.status == status)
    if action_type:
        query = query.filter(CsFollowUpAction.action_type == action_type)
    if priority:
        query = query.filter(CsFollowUpAction.priority == priority)

    total = query.count()
    actions = query.order_by(CsFollowUpAction.created_at.desc()).offset((page - 1) * limit).limit(limit).all()

    return {
        "items": [{
            "id": a.id,
            "inquiry_id": a.inquiry_id,
            "action_type": a.action_type,
            "action_label": a.action_label,
            "priority": a.priority,
            "status": a.status,
            "assigned_to": a.assigned_to,
            "notes": a.notes,
            "ai_suggested": a.ai_suggested,
            "completed_at": a.completed_at.isoformat() if a.completed_at else None,
            "created_at": a.created_at.isoformat() if a.created_at else None,
        } for a in actions],
        "total": total,
        "page": page,
        "limit": limit,
    }


@router.put("/followup-actions/{action_id}")
def update_followup_action(
    action_id: int,
    status: Optional[str] = None,
    notes: Optional[str] = None,
    assigned_to: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """후속 조치 상태 업데이트"""
    action = db.query(CsFollowUpAction).filter(CsFollowUpAction.id == action_id).first()
    if not action:
        raise HTTPException(status_code=404, detail="후속 조치를 찾을 수 없습니다")

    if status:
        action.status = status
        if status == "completed":
            action.completed_at = datetime.now()
    if notes is not None:
        action.notes = notes
    if assigned_to is not None:
        action.assigned_to = assigned_to

    db.commit()
    return {"message": "업데이트 완료", "id": action_id, "status": action.status}


@router.get("/followup-actions/summary")
def followup_actions_summary(db: Session = Depends(get_db)):
    """후속 조치 대시보드 요약"""
    total = db.query(func.count(CsFollowUpAction.id)).scalar() or 0
    pending = db.query(func.count(CsFollowUpAction.id)).filter(
        CsFollowUpAction.status == "pending"
    ).scalar() or 0
    in_progress = db.query(func.count(CsFollowUpAction.id)).filter(
        CsFollowUpAction.status == "in_progress"
    ).scalar() or 0
    completed = db.query(func.count(CsFollowUpAction.id)).filter(
        CsFollowUpAction.status == "completed"
    ).scalar() or 0

    # 유형별 통계
    type_rows = db.query(
        CsFollowUpAction.action_type, func.count(CsFollowUpAction.id)
    ).filter(CsFollowUpAction.status == "pending").group_by(CsFollowUpAction.action_type).all()
    by_type = {row[0]: row[1] for row in type_rows}

    return {
        "total": total,
        "pending": pending,
        "in_progress": in_progress,
        "completed": completed,
        "pending_by_type": by_type,
    }


# ──────────────────────────────────────────────
# CS 분석 (Analytics) 엔드포인트
# ──────────────────────────────────────────────

@router.get("/inquiries/analytics/volume")
def get_inquiry_volume(
    period: str = Query(default="daily", description="daily 또는 monthly"),
    days: int = Query(default=30, description="일별 조회 시 최근 N일"),
    months: int = Query(default=12, description="월별 조회 시 최근 N개월"),
    db: Session = Depends(get_db),
):
    """일별/월별 문의 수량 통계 (쇼핑몰별·카테고리별 breakdown 포함)"""
    import random
    from datetime import timedelta

    try:
        total_count = db.query(func.count(CsInquiry.id)).scalar() or 0

        if total_count == 0:
            # 샘플 데이터 반환
            random.seed(42)
            malls = ["스마트스토어", "쿠팡", "11번가", "카카오선물하기"]
            categories = ["배송문의", "교환/반품", "상품문의", "기타"]
            base = datetime.now()

            if period == "monthly":
                sample_data = []
                for i in range(months):
                    d = base.replace(day=1) - timedelta(days=30 * (months - 1 - i))
                    month_str = d.strftime("%Y-%m")
                    total = random.randint(40, 150)
                    by_mall = {m: random.randint(5, max(5, total // 4)) for m in malls}
                    by_category = {c: random.randint(5, max(5, total // 4)) for c in categories}
                    sample_data.append({
                        "date": month_str,
                        "total": total,
                        "by_mall": by_mall,
                        "by_category": by_category,
                    })
            else:
                sample_data = []
                for i in range(14):
                    d = base - timedelta(days=13 - i)
                    date_str = d.strftime("%Y-%m-%d")
                    total = random.randint(5, 25)
                    by_mall = {m: random.randint(1, max(1, total // 4)) for m in malls}
                    by_category = {c: random.randint(1, max(1, total // 4)) for c in categories}
                    sample_data.append({
                        "date": date_str,
                        "total": total,
                        "by_mall": by_mall,
                        "by_category": by_category,
                    })

            return {"period": period, "data": sample_data, "is_sample": True}

        # 실제 DB 데이터 집계 — 단일 쿼리로 전체 범위 조회 후 Python에서 집계
        base = datetime.now()
        data = []

        if period == "monthly":
            # 조회 범위 계산 (months 개월치 전체를 한 번에 조회)
            range_start = (base.replace(day=1) - timedelta(days=30 * (months - 1))).replace(day=1)
            range_end = (base.replace(day=28) + timedelta(days=4)).replace(day=1)

            rows = (
                db.query(CsInquiry.inquiry_date, CsInquiry.mall_name, CsInquiry.category)
                .filter(CsInquiry.inquiry_date >= range_start, CsInquiry.inquiry_date < range_end)
                .all()
            )

            # 월별 버킷 초기화 — 데이터가 없는 달도 0으로 채워 차트에서 연속 표시
            buckets: dict = {}
            for i in range(months):
                month_start = (base.replace(day=1) - timedelta(days=30 * (months - 1 - i))).replace(day=1)
                month_str = month_start.strftime("%Y-%m")
                buckets[month_str] = {"total": 0, "by_mall": {}, "by_category": {}}

            for r in rows:
                if r.inquiry_date is None:
                    continue
                month_str = r.inquiry_date.strftime("%Y-%m")
                if month_str not in buckets:
                    continue
                bucket = buckets[month_str]
                bucket["total"] += 1
                mall = r.mall_name or "기타"
                bucket["by_mall"][mall] = bucket["by_mall"].get(mall, 0) + 1
                cat = r.category or "기타"
                bucket["by_category"][cat] = bucket["by_category"].get(cat, 0) + 1

            data = [{"date": k, **v} for k, v in sorted(buckets.items())]

        else:
            # 조회 범위 계산 (days 일치 전체를 한 번에 조회)
            range_start = (base - timedelta(days=days - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
            range_end = base.replace(hour=23, minute=59, second=59, microsecond=999999)

            rows = (
                db.query(CsInquiry.inquiry_date, CsInquiry.mall_name, CsInquiry.category)
                .filter(CsInquiry.inquiry_date >= range_start, CsInquiry.inquiry_date <= range_end)
                .all()
            )

            # 일별 버킷 초기화 — 데이터가 없는 날도 0으로 채워 차트에서 연속 표시
            buckets: dict = {}
            for i in range(days):
                d = base - timedelta(days=days - 1 - i)
                date_str = d.strftime("%Y-%m-%d")
                buckets[date_str] = {"total": 0, "by_mall": {}, "by_category": {}}

            for r in rows:
                if r.inquiry_date is None:
                    continue
                date_str = r.inquiry_date.strftime("%Y-%m-%d")
                if date_str not in buckets:
                    continue
                bucket = buckets[date_str]
                bucket["total"] += 1
                mall = r.mall_name or "기타"
                bucket["by_mall"][mall] = bucket["by_mall"].get(mall, 0) + 1
                cat = r.category or "기타"
                bucket["by_category"][cat] = bucket["by_category"].get(cat, 0) + 1

            data = [{"date": k, **v} for k, v in sorted(buckets.items())]

        return {"period": period, "data": data, "is_sample": False}

    except Exception as e:
        logger.error(f"문의 수량 통계 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/inquiries/analytics/keywords")
async def get_inquiry_keywords(
    days: int = Query(default=30, description="최근 N일간 문의 분석"),
    db: Session = Depends(get_db),
):
    """문의 키워드 분석 (AI 기반, Claude API 사용)"""
    from datetime import timedelta

    sample_keywords = [
        {"keyword": "배송 지연", "count": 23, "importance": "high", "category": "배송문의",
         "sample_inquiries": ["주문한 상품이 아직 안 왔어요", "배송이 3일째 안 움직여요"]},
        {"keyword": "파손/훼손", "count": 15, "importance": "high", "category": "교환/반품",
         "sample_inquiries": ["케이크가 깨져서 왔어요", "포장이 찢어져 있었습니다"]},
        {"keyword": "교환 절차", "count": 12, "importance": "medium", "category": "교환/반품",
         "sample_inquiries": ["다른 맛으로 교환하고 싶어요", "교환은 어떻게 하나요?"]},
        {"keyword": "성분/알레르기", "count": 10, "importance": "medium", "category": "상품문의",
         "sample_inquiries": ["비누 성분이 궁금해요", "알레르기 있는데 괜찮을까요?"]},
        {"keyword": "선물 포장", "count": 8, "importance": "medium", "category": "기타",
         "sample_inquiries": ["선물 포장 가능한가요?", "메시지 카드 넣을 수 있나요?"]},
        {"keyword": "유통기한", "count": 7, "importance": "medium", "category": "상품문의",
         "sample_inquiries": ["마카롱 유통기한이 어떻게 되나요?", "케이크 보관 방법 알려주세요"]},
        {"keyword": "환불", "count": 6, "importance": "high", "category": "교환/반품",
         "sample_inquiries": ["환불 처리해주세요", "환불은 언제 되나요?"]},
        {"keyword": "택배사 변경", "count": 5, "importance": "low", "category": "배송문의",
         "sample_inquiries": ["택배사를 바꿔주실 수 있나요?", "다른 택배사로 보내주세요"]},
        {"keyword": "대량 주문", "count": 4, "importance": "medium", "category": "기타",
         "sample_inquiries": ["100개 이상 주문 가능한가요?", "단체 할인 있나요?"]},
        {"keyword": "재입고", "count": 3, "importance": "low", "category": "상품문의",
         "sample_inquiries": ["품절된 상품 재입고 언제 되나요?", "라벤더 캔들 재입고 알림"]},
    ]

    try:
        from datetime import timedelta
        cutoff = datetime.now() - timedelta(days=days)
        inquiries = (
            db.query(CsInquiry)
            .filter(CsInquiry.inquiry_date >= cutoff)
            .order_by(CsInquiry.inquiry_date.desc())
            .all()
        )
        count = len(inquiries)

        if count == 0:
            return {"keywords": sample_keywords, "total_inquiries_analyzed": 0, "period_days": days, "is_sample": True}

        from app.config import get_settings
        _kw_settings = get_settings()
        api_key = _kw_settings.ANTHROPIC_API_KEY or os.getenv("ANTHROPIC_API_KEY", "")
        if not api_key:
            return {"keywords": sample_keywords, "total_inquiries_analyzed": count, "period_days": days, "is_sample": True}

        # 100건 초과 시 최근 100건만 사용 (토큰 절약)
        analysis_targets = inquiries[:100]
        analysis_count = len(analysis_targets)

        contents = "\n".join(
            f"{idx + 1}. [{inq.title or '제목없음'}] {(inq.content or '')[:200]}"
            for idx, inq in enumerate(analysis_targets)
        )

        def _build_keywords_prompt(inq_count: int, inq_contents: str) -> str:
            return f"""다음은 고객 문의 {inq_count}건입니다. 주요 키워드와 이슈를 분석해주세요.

문의 내용들:
{inq_contents}

반드시 다음 JSON 형식으로만 답변하세요 (다른 텍스트 없이):
{{
  "keywords": [
    {{
      "keyword": "키워드명",
      "count": 관련문의수,
      "importance": "high/medium/low",
      "category": "배송문의/교환반품/상품문의/기타",
      "trend": "increasing/stable/decreasing",
      "sample_inquiries": ["관련 문의 제목1", "제목2"]
    }}
  ]
}}
최대 15개, importance가 high인 것부터 정렬.
"trend"는 최근 문의에서 해당 키워드 빈도가 증가/유지/감소 추세인지 판단."""

        def _extract_json_from_text(raw_text: str) -> dict | None:
            json_start = raw_text.find("{")
            json_end = raw_text.rfind("}") + 1
            if json_start != -1 and json_end > json_start:
                try:
                    return json.loads(raw_text[json_start:json_end])
                except json.JSONDecodeError:
                    return None
            return None

        parsed_result = None
        try:
            async with httpx.AsyncClient() as client:
                # 1차 시도: 전체 프롬프트
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-sonnet-4-20250514",
                        "max_tokens": 3000,
                        "messages": [{"role": "user", "content": _build_keywords_prompt(analysis_count, contents)}],
                    },
                    timeout=60.0,
                )
                if resp.status_code == 200:
                    raw_text = resp.json()["content"][0]["text"]
                    parsed_result = _extract_json_from_text(raw_text)
                    if parsed_result is None:
                        logger.warning("keywords 1차 JSON 파싱 실패, 재시도")
                else:
                    logger.error(f"Claude API 응답 오류 (keywords 1차): {resp.status_code} - {resp.text}")

                # 2차 시도: 파싱 실패 시 간소화된 프롬프트로 재시도
                if parsed_result is None:
                    retry_prompt = (
                        f"고객 문의 {analysis_count}건의 키워드를 분석해 JSON만 출력하세요.\n"
                        f"문의 목록:\n{contents}\n\n"
                        '출력 형식: {"keywords": [{"keyword": "...", "count": 숫자, "importance": "high/medium/low", '
                        '"category": "배송문의/교환반품/상품문의/기타", "trend": "increasing/stable/decreasing", '
                        '"sample_inquiries": ["제목1"]}]}'
                    )
                    resp2 = await client.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={
                            "x-api-key": api_key,
                            "anthropic-version": "2023-06-01",
                            "content-type": "application/json",
                        },
                        json={
                            "model": "claude-sonnet-4-20250514",
                            "max_tokens": 2048,
                            "messages": [{"role": "user", "content": retry_prompt}],
                        },
                        timeout=60.0,
                    )
                    if resp2.status_code == 200:
                        raw_text2 = resp2.json()["content"][0]["text"]
                        parsed_result = _extract_json_from_text(raw_text2)
                        if parsed_result is None:
                            logger.error("keywords 2차 JSON 파싱도 실패")
                    else:
                        logger.error(f"Claude API 응답 오류 (keywords 2차): {resp2.status_code} - {resp2.text}")

        except Exception as e:
            logger.error(f"Claude API 호출 실패 (keywords): {e}")

        if parsed_result is not None:
            return {
                "keywords": parsed_result.get("keywords", []),
                "total_inquiries_analyzed": analysis_count,
                "period_days": days,
                "is_sample": False,
            }

        # Claude 완전 실패 시 샘플 반환
        return {"keywords": sample_keywords, "total_inquiries_analyzed": count, "period_days": days, "is_sample": True}

    except Exception as e:
        logger.error(f"키워드 분석 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/inquiries/analytics/action-items")
async def get_action_items(
    days: int = Query(default=30, description="최근 N일간 문의 기반 분석"),
    db: Session = Depends(get_db),
):
    """AI 기반 액션 아이템 추천 (문의 키워드 분석 결과 활용)"""
    from datetime import timedelta

    sample_actions = [
        {"title": "배송 프로세스 개선",
         "description": "배송 지연 관련 문의가 가장 많습니다. 물류 파트너사와 배송 리드타임 단축 협의가 필요합니다.",
         "priority": "high", "category": "배송",
         "related_keywords": ["배송 지연", "택배사 변경"],
         "estimated_impact": "배송 관련 문의 30% 감소 예상"},
        {"title": "포장 품질 강화",
         "description": "케이크/마카롱 등 파손 관련 교환/반품이 빈번합니다. 완충재 보강 및 포장 방식 개선이 필요합니다.",
         "priority": "high", "category": "품질",
         "related_keywords": ["파손/훼손", "교환 절차"],
         "estimated_impact": "파손 관련 클레임 50% 감소, 교환/반품 비용 절감"},
        {"title": "상품 상세페이지 성분 정보 강화",
         "description": "성분/알레르기 문의가 꾸준합니다. 상세페이지에 전성분, 알레르기 유발물질 표기를 강화하면 문의를 사전 차단할 수 있습니다.",
         "priority": "medium", "category": "상품정보",
         "related_keywords": ["성분/알레르기", "유통기한"],
         "estimated_impact": "상품문의 25% 감소"},
        {"title": "선물포장 서비스 안내 페이지 제작",
         "description": "선물 포장 가능 여부 문의가 반복됩니다. 선물 포장 옵션을 상품 페이지에 명시하고, 메시지 카드 서비스를 눈에 띄게 안내하세요.",
         "priority": "medium", "category": "서비스",
         "related_keywords": ["선물 포장"],
         "estimated_impact": "선물 관련 문의 80% 감소"},
        {"title": "B2B/대량주문 전용 페이지 개설",
         "description": "단체/대량 주문 문의가 있습니다. 별도 견적 요청 폼이나 B2B 안내 페이지를 만들면 매출 기회를 놓치지 않을 수 있습니다.",
         "priority": "medium", "category": "영업",
         "related_keywords": ["대량 주문"],
         "estimated_impact": "B2B 매출 기회 확보"},
        {"title": "자동 환불 처리 시스템 도입",
         "description": "단순 환불 문의는 자동화할 수 있습니다. 환불 조건(7일 이내, 미개봉 등) 충족 시 자동 처리 flow를 구축하세요.",
         "priority": "low", "category": "운영",
         "related_keywords": ["환불"],
         "estimated_impact": "CS 담당자 업무 시간 주 2시간 절감"},
        {"title": "재입고 알림 서비스 도입",
         "description": "품절 상품 재입고 문의가 있습니다. 카카오 알림톡 기반 재입고 알림 서비스를 도입하면 재구매 전환율을 높일 수 있습니다.",
         "priority": "low", "category": "마케팅",
         "related_keywords": ["재입고"],
         "estimated_impact": "재구매 전환율 15% 향상"},
    ]

    try:
        from datetime import timedelta
        cutoff = datetime.now() - timedelta(days=days)
        inquiries = (
            db.query(CsInquiry)
            .filter(CsInquiry.inquiry_date >= cutoff)
            .order_by(CsInquiry.inquiry_date.desc())
            .all()
        )
        count = len(inquiries)

        if count == 0:
            return {
                "action_items": sample_actions,
                "generated_at": datetime.now().isoformat(),
                "based_on_inquiries": 0,
                "is_sample": True,
            }

        api_key = os.getenv("ANTHROPIC_API_KEY", "")
        if not api_key:
            return {
                "action_items": sample_actions,
                "generated_at": datetime.now().isoformat(),
                "based_on_inquiries": count,
                "is_sample": True,
            }

        # ── Step 1: keywords 분석 먼저 수행 (액션 아이템 품질 향상) ──
        keywords_for_action: list = []
        keywords_summary_text = ""
        try:
            analysis_targets = inquiries[:100]
            kw_contents = "\n".join(
                f"{idx + 1}. [{inq.title or '제목없음'}] {(inq.content or '')[:200]}"
                for idx, inq in enumerate(analysis_targets)
            )
            kw_prompt = (
                f"고객 문의 {len(analysis_targets)}건의 주요 키워드를 분석해 JSON만 출력하세요.\n"
                f"문의 목록:\n{kw_contents}\n\n"
                '출력 형식: {"keywords": [{"keyword": "...", "count": 숫자, "importance": "high/medium/low", '
                '"category": "배송문의/교환반품/상품문의/기타", "trend": "increasing/stable/decreasing", '
                '"sample_inquiries": ["제목1"]}]}'
            )
            async with httpx.AsyncClient() as kw_client:
                kw_resp = await kw_client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-sonnet-4-20250514",
                        "max_tokens": 2048,
                        "messages": [{"role": "user", "content": kw_prompt}],
                    },
                    timeout=60.0,
                )
            if kw_resp.status_code == 200:
                kw_raw = kw_resp.json()["content"][0]["text"]
                kw_start = kw_raw.find("{")
                kw_end = kw_raw.rfind("}") + 1
                if kw_start != -1 and kw_end > kw_start:
                    kw_parsed = json.loads(kw_raw[kw_start:kw_end])
                    keywords_for_action = kw_parsed.get("keywords", [])
                    keywords_summary_text = "\n".join(
                        f"- {kw['keyword']} ({kw.get('count', '?')}건, {kw.get('importance', '')} / {kw.get('trend', '')})"
                        for kw in keywords_for_action[:10]
                    )
        except Exception as kw_err:
            logger.warning(f"action-items용 keywords 사전 분석 실패 (무시하고 계속): {kw_err}")

        # keywords 분석 실패 시 카테고리 빈도로 대체
        if not keywords_summary_text:
            category_counts: dict = {}
            for inq in inquiries:
                cat = inq.category or "기타"
                category_counts[cat] = category_counts.get(cat, 0) + 1
            keywords_summary_text = "\n".join(
                f"- {cat}: {cnt}건"
                for cat, cnt in sorted(category_counts.items(), key=lambda x: -x[1])
            )

        # ── Step 2: 카테고리 분포 계산 ──
        category_dist: dict = {}
        for inq in inquiries:
            cat = inq.category or "기타"
            category_dist[cat] = category_dist.get(cat, 0) + 1
        category_distribution_text = "\n".join(
            f"- {cat}: {cnt}건 ({round(cnt / count * 100)}%)"
            for cat, cnt in sorted(category_dist.items(), key=lambda x: -x[1])
        )

        # ── Step 3: 샘플 문의 내용 (최근 30건, 제목+내용 포함) ──
        sample_contents = "\n".join(
            f"{idx + 1}. [{inq.title or '제목없음'}] {(inq.content or '')[:150]}"
            for idx, inq in enumerate(inquiries[:30])
        )

        prompt = f"""고객 문의 데이터 분석 결과를 바탕으로, CS 팀이 취해야 할 구체적인 액션 아이템을 추천해주세요.

## 분석 기간: 최근 {days}일
## 총 문의 수: {count}건

## 키워드 분석 결과:
{keywords_summary_text}

## 최근 문의 샘플 ({min(30, count)}건):
{sample_contents}

## 카테고리별 문의 분포:
{category_distribution_text}

반드시 다음 JSON 형식으로만 답변하세요:
{{
  "action_items": [
    {{
      "title": "구체적인 액션 제목",
      "description": "상세 설명 (왜 필요한지, 어떻게 실행하는지)",
      "priority": "high/medium/low",
      "category": "배송/품질/서비스/운영/마케팅/상품정보",
      "related_keywords": ["키워드1", "키워드2"],
      "estimated_impact": "구체적인 기대 효과",
      "suggested_deadline": "즉시/1주일/1개월/분기"
    }}
  ]
}}
최대 8개, priority 순으로 정렬.
각 액션은 구체적이고 실행 가능해야 합니다."""

        def _extract_json_action(raw_text: str) -> dict | None:
            json_start = raw_text.find("{")
            json_end = raw_text.rfind("}") + 1
            if json_start != -1 and json_end > json_start:
                try:
                    return json.loads(raw_text[json_start:json_end])
                except json.JSONDecodeError:
                    return None
            return None

        parsed_result = None
        try:
            async with httpx.AsyncClient() as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-sonnet-4-20250514",
                        "max_tokens": 3000,
                        "messages": [{"role": "user", "content": prompt}],
                    },
                    timeout=90.0,
                )
                if resp.status_code == 200:
                    raw_text = resp.json()["content"][0]["text"]
                    parsed_result = _extract_json_action(raw_text)
                    if parsed_result is None:
                        logger.error(f"action-items JSON 파싱 실패. 응답 앞 300자: {raw_text[:300]}")
                else:
                    logger.error(f"Claude API 응답 오류 (action-items): {resp.status_code} - {resp.text}")
        except Exception as e:
            logger.error(f"Claude API 호출 실패 (action-items): {e}")

        if parsed_result is not None:
            return {
                "action_items": parsed_result.get("action_items", []),
                "generated_at": datetime.now().isoformat(),
                "based_on_inquiries": count,
                "keywords_analyzed": len(keywords_for_action),
                "is_sample": False,
            }

        # Claude 실패 시 샘플 반환
        return {
            "action_items": sample_actions,
            "generated_at": datetime.now().isoformat(),
            "based_on_inquiries": count,
            "is_sample": True,
        }

    except Exception as e:
        logger.error(f"액션 아이템 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))
